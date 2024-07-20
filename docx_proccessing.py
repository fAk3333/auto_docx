import os
import shutil
from docx2pdf import convert
from pdf2image import convert_from_path
from docx.shared import Inches
import docx
import docx2python


def folder_creater(name_of_project):
    try:
        os.makedirs(name_of_project)
    except:
        print('дир уже создан')
    ary_of_files = []
    for root, dirs, files in os.walk("templates"):
        for filename in files:
            ary_of_files.append(filename)
    print(ary_of_files)
    for file in ary_of_files:
        src = f'templates/{file}'
        dest = f'{name_of_project}/{name_of_project} {file}'
        shutil.copy2(src, dest)

def word_processor(name_of_project, cipher, customer, name_of_obj, leng):
    old_cipher = '32211846356-1-1321-ИГДИ'
    old_customer = 'ООО «ТЕХИННОВАЦИЯ»'
    old_name_of_obj = 'Строительство газопровода к жилому дому, расположенному по адресу: г. ' \
                      'Севастополь, с. Осипенко, ул. Ветеранов, д. 34 ' \
                      '(кадастровый номер земельного участка 91:04:036001:1321). ' \
                      'Инженерные изыскания'
    old_leng = '0,15'
    for root, dirs, files in os.walk(name_of_project):
        for filename in files:
            doc = docx.Document(f'{name_of_project}/{filename}')
            for paragraph in doc.paragraphs:
                if old_cipher in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_cipher, cipher)
                if old_customer in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_customer, customer)
                if old_name_of_obj in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_name_of_obj, name_of_obj)
                if old_leng in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_leng, leng)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if old_cipher in paragraph.text:
                                paragraph.text = paragraph.text.replace(old_cipher, cipher)
                            if old_customer in paragraph.text:
                                paragraph.text = paragraph.text.replace(old_customer, customer)
                            if old_name_of_obj in paragraph.text:
                                paragraph.text = paragraph.text.replace(old_name_of_obj, name_of_obj)
                            if old_leng in paragraph.text:
                                paragraph.text = paragraph.text.replace(old_leng, leng)
            doc.save(f'{name_of_project}/{filename}')
            print('Замены готовы!')


def docx_to_pdf_to_jpeg(name_of_project):
    for root, dirs, files in os.walk(name_of_project):
        for filename in files:
            old = filename
            new = filename.replace('.docx', '.pdf')
            convert(f'{name_of_project}/{old}', f'{name_of_project}/{new}')
    pdf_ary = []
    for root, dirs, files in os.walk(name_of_project):
        for filename in files:
            if '.pdf' in filename:
                pdf_ary.append(filename)
    for pdf in pdf_ary:
        pages = convert_from_path(f'{name_of_project}/{pdf}', poppler_path=r'poppler-0.68.0\bin')
        for i in range(len(pages)):
            pages[i].save(name_of_project + '/' + pdf.replace('.pdf', '') + 'page' + str(i) + '.jpg', 'JPEG')
    print('Готово')